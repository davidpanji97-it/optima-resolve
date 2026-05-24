import os
#import sentry_sdk
#sentry_sdk.init(
    #dsn="MASUKKAN_LINK_DSN_DARI_SENTRY_DI_SINI",
    # Setel ke 1.0 untuk merekam 100% error yang terjadi
    #traces_sample_rate=1.0,
#)
import io
import time
import json
from datetime import datetime, timedelta

import streamlit as st
import pandas as pd
import plotly.express as px
from dotenv import load_dotenv

# Import library pembuat Word
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Konfigurasi Firebase NoSQL
import firebase_admin
from firebase_admin import credentials, firestore

# Konfigurasi AI & RAG
from groq import Groq
from langchain_community.document_loaders import TextLoader
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import CharacterTextSplitter

# 1. KONFIGURASI HALAMAN UTAMA (TIDAK ADA CSS PAKSAAN)
st.set_page_config(page_title="Optima Resolve", page_icon="💡", layout="wide", initial_sidebar_state="collapsed")

load_dotenv()
api_key = os.getenv("GROQ_API_KEY")

# 2. INIT FIREBASE
if not firebase_admin._apps:
    try:
        if os.path.exists("credentials-firebase.json"):
            cred = credentials.Certificate("credentials-firebase.json")
        else:
            firebase_secrets = json.loads(st.secrets["firebase_json"])
            cred = credentials.Certificate(firebase_secrets)
            
        firebase_admin.initialize_app(cred)
    except Exception as e:
        st.error(f"Gagal memuat kunci Firebase. Error: {e}")

db = firestore.client()

# 3. PERSIAPAN AI KNOWLEDGE BASE
@st.cache_resource
def prepare_knowledge_base():
    if not os.path.exists("manual_it.txt"):
        with open("manual_it.txt", "w") as f:
            f.write("Panduan Dasar IT Optima Resolve. Hubungi tim teknis jika masalah berlanjut.")
            
    loader = TextLoader("manual_it.txt")
    documents = loader.load()
    text_splitter = CharacterTextSplitter(chunk_size=600, chunk_overlap=100)
    docs = text_splitter.split_documents(documents)
    embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    return FAISS.from_documents(docs, embeddings)

try:
    vector_db = prepare_knowledge_base()
except:
    vector_db = None

# 4. FUNGSI PEMBANTU
def generate_sequential_id():
    prefix = "OPT-"
    try:
        tickets_ref = db.collection('rekap_tiket')
        query = tickets_ref.order_by('Waktu', direction=firestore.Query.DESCENDING).limit(1).stream()
        last_id_string = None
        for doc in query:
            last_id_string = doc.to_dict().get('ID Tiket')
            
        if last_id_string:
            last_number = int(last_id_string.replace(prefix, ""))
            next_number = last_number + 1
        else:
            next_number = 1
    except Exception:
        next_number = 1
    return f"{prefix}{next_number:05d}"

def logout():
    st.session_state.clear()
    st.rerun()

def create_word_docx(info, tujuan, ruang, pihak, langkah):
    doc = Document()
    doc.add_heading('STANDARD OPERATING PROCEDURE', 0)
    doc.add_heading(info.upper(), level=1)
    doc.add_heading('1. TUJUAN', level=2)
    doc.add_paragraph(tujuan if tujuan else '-')
    doc.add_heading('2. RUANG LINGKUP', level=2)
    doc.add_paragraph(ruang if ruang else '-')
    doc.add_heading('3. PIHAK TERKAIT', level=2)
    doc.add_paragraph(pihak if pihak else '-')
    doc.add_heading('4. LANGKAH-LANGKAH PROSEDUR', level=2)
    for l in langkah.split('\n'):
        if l.strip() != "":
            doc.add_paragraph(l.strip(), style='List Bullet')
    doc.add_paragraph(f"\n\nDibuat pada: {datetime.now().strftime('%d %B %Y')}")
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# 5. STATE MANAGEMENT
if "page" not in st.session_state: st.session_state.page = "HOME"
if "target_page" not in st.session_state: st.session_state.target_page = "FORM_TIKET" 
if "user_data" not in st.session_state: st.session_state.user_data = {}
if "ticket_id" not in st.session_state: st.session_state.ticket_id = generate_sequential_id()
if "messages" not in st.session_state: st.session_state.messages = []
if "ticket_status" not in st.session_state: st.session_state.ticket_status = "Open ⏳"
if "ticket_priority" not in st.session_state: st.session_state.ticket_priority = "Belum Terdeteksi"
if "uploader_key" not in st.session_state: st.session_state.uploader_key = 0
if "first_prompt_pending" not in st.session_state: st.session_state.first_prompt_pending = False

DATABASE_KARYAWAN = {
    "12345": {"pass": "yola123", "nama": "Yola Suryani", "divisi": "HR", "telepon": "0812-333-444"},
    "123456": {"pass": "testing123", "nama": "Public User", "divisi": "Finance", "telepon": "0856-777-888"},
    "11223": {"pass": "honda123", "nama": "Andi Pratama", "divisi": "IT", "telepon": "0811-222-333"} # AKUN KHUSUS IT SUPPORT
}

# ==========================================
# ROUTER 1: HOME PAGE
# ==========================================
if st.session_state.page == "HOME":
    st.title("Welcome to Optima Resolve")
    
    if "nip" in st.session_state.user_data:
        st.subheader(f"Kembali bekerja, {st.session_state.user_data['nama']}!")
    
    # DESKRIPSI BARU SESUAI REQUEST
    st.markdown("Sistem asisten pintar yang siap membimbing Anda menyelesaikan kendala teknis secara mandiri, sekaligus mengelola administrasi personal Anda dengan cepat dan akurat.")
    st.divider()
    
    # FITUR KHUSUS IT SUPPORT (HANYA MUNCUL JIKA LOGIN SEBAGAI DIVISI IT)
    if st.session_state.user_data.get('divisi') == 'IT':
        st.success("👨‍💻 **PANEL IT SUPPORT (KONTROL TIKET)**")
        with st.container(border=True):
            col_id, col_stat, col_btn = st.columns([2, 2, 1])
            with col_id:
                # Mengambil semua ID tiket yang ada di database untuk dipilih IT Support
                try:
                    all_tickets = [doc.to_dict().get('ID Tiket') for doc in db.collection('rekap_tiket').stream()]
                except:
                    all_tickets = []
                pilih_tiket = st.selectbox("Pilih ID Tiket User:", all_tickets if all_tickets else ["Tidak ada tiket"])
            with col_stat:
                status_baru = st.selectbox("Ubah Status Menjadi:", ["Progress 🔄", "Resolved ✅", "Pending ⏳"])
            with col_btn:
                st.write("")
                st.write("")
                if st.button("Update Status", type="primary", use_container_width=True):
                    if pilih_tiket != "Tidak ada tiket":
                        db.collection('rekap_tiket').document(pilih_tiket).update({'Status': status_baru})
                        st.toast(f"Status tiket {pilih_tiket} berhasil diubah!")
        st.divider()

    colA, colB = st.columns(2)
    with colA:
        with st.container(border=True):
            st.subheader("🛠️ IT Helpdesk")
            st.write("Laporkan kendala perangkat keras, jaringan, atau sistem IT. Didukung oleh AI Scanner.")
            if st.button("Buat Tiket Kendala IT", use_container_width=True, type="primary"):
                if "nip" in st.session_state.user_data:
                    st.session_state.ticket_id = generate_sequential_id() # Reset ID tiket baru untuk user
                    st.session_state.messages = [] # Bersihkan chat lama
                    st.session_state.page = "FORM_TIKET"
                else:
                    st.session_state.target_page = "FORM_TIKET"
                    st.session_state.page = "LOGIN"
                st.rerun()

    with colB:
        with st.container(border=True):
            st.subheader("📄 SOP Generator")
            st.write("Buat contoh Standard Operating Procedure (SOP) secara instan dan berstandar resmi.")
            if st.button("Buat Dokumen SOP", use_container_width=True, type="primary"):
                if "nip" in st.session_state.user_data:
                    st.session_state.page = "SOP_MODULE"
                else:
                    st.session_state.target_page = "SOP_MODULE"
                    st.session_state.page = "LOGIN"
                st.rerun()
            
    if "nip" in st.session_state.user_data:
        st.divider()
        if st.button("🚪 Keluar Sistem (Logout)"): 
            logout()

# ==========================================
# ROUTER 2: LOGIN
# ==========================================
elif st.session_state.page == "LOGIN":
    col1, col2, col3 = st.columns([1, 1.5, 1])
    with col2:
        with st.form("login_form", border=True):
            st.subheader("💡 Login Optima Resolve")
            
            nip_input = st.text_input("NIP / Employee ID")
            pass_input = st.text_input("Password", type="password")
            
            submitted = st.form_submit_button("Masuk", use_container_width=True, type="primary")
            
            if submitted:
                if nip_input == "Admin" and pass_input == "william123":
                    st.session_state.page = "ADMIN"
                    st.rerun()
                elif nip_input in DATABASE_KARYAWAN and DATABASE_KARYAWAN[nip_input]["pass"] == pass_input:
                    st.session_state.user_data = DATABASE_KARYAWAN[nip_input]
                    st.session_state.user_data['nip'] = nip_input
                    st.session_state.page = st.session_state.target_page
                    st.rerun()
                else:
                    st.error("❌ Login Gagal. NIP atau Password tidak terdaftar.")
                    
        if st.button("⬅️ Kembali ke Beranda", use_container_width=True):
            st.session_state.page = "HOME"
            st.rerun()

# ==========================================
# ROUTER 3: SOP MODULE
# ==========================================
elif st.session_state.page == "SOP_MODULE":
    st.title("📄 Generator SOP Internal")
    
    col_form, col_result = st.columns([1, 1.2])
    with col_form:
        with st.form("form_sop", border=True):
            info_dokumen = st.text_input("Informasi Dokumen")
            tujuan = st.text_area("Tujuan")
            ruang_lingkup = st.text_area("Ruang Lingkup")
            pihak_terkait = st.text_input("Pihak Terkait")
            langkah_langkah = st.text_area("Langkah-langkah Prosedur", height=150)
            
            submit_sop = st.form_submit_button("Preview SOP 🚀", type="primary")
            
            if submit_sop:
                if info_dokumen == "" or langkah_langkah == "":
                    st.warning("⚠️ Informasi dan Langkah-langkah wajib diisi!")
                else:
                    st.session_state.generated_sop = {"info": info_dokumen, "tujuan": tujuan, "ruang": ruang_lingkup, "pihak": pihak_terkait, "langkah": langkah_langkah}

    with col_result:
        if "generated_sop" in st.session_state:
            sop = st.session_state.generated_sop
            with st.container(border=True):
                st.success("✅ Dokumen SOP Berhasil Dibuat!")
                if DOCX_AVAILABLE:
                    word_file = create_word_docx(sop["info"], sop["tujuan"], sop["ruang"], sop["pihak"], sop["langkah"])
                    st.download_button(label="📥 Download format Word (.docx)", data=word_file, file_name=f"SOP_{sop['info'].replace(' ', '_')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary", use_container_width=True)

                st.subheader("STANDARD OPERATING PROCEDURE")
                st.write(f"**{sop['info'].upper()}**")
                st.divider()
                st.write("**1. TUJUAN**\n" + (sop['tujuan'] if sop['tujuan'] else '-'))
                st.write("**2. RUANG LINGKUP**\n" + (sop['ruang'] if sop['ruang'] else '-'))
                st.write("**3. PIHAK TERKAIT**\n" + (sop['pihak'] if sop['pihak'] else '-'))
                st.write("**4. LANGKAH-LANGKAH PROSEDUR**")
                for langkah in sop['langkah'].split('\n'):
                    if langkah.strip() != "": st.write(f"- {langkah}")
        else:
            st.info("Isi formulir di sebelah kiri dan klik 'Preview SOP'.")

    st.divider()
    if st.button("⬅️ Kembali ke Beranda"):
        if "generated_sop" in st.session_state: del st.session_state.generated_sop
        st.session_state.page = "HOME"
        st.rerun()

# ==========================================
# ROUTER 4: FORM TIKET
# ==========================================
elif st.session_state.page == "FORM_TIKET":
    st.title("📝 Buat Tiket Kendala Baru")
    
    with st.form("form_keluhan", border=True):
        st.write("**Detail Pelapor (Sistem HRD)**")
        colA, colB = st.columns(2)
        with colA:
            nama = st.text_input("Nama Lengkap", value=st.session_state.user_data.get('nama', ''), disabled=True)
            divisi = st.text_input("Divisi", value=st.session_state.user_data.get('divisi', ''), disabled=True)
        with colB:
            telepon = st.text_input("Nomor Telepon", value=st.session_state.user_data.get('telepon', ''), disabled=True)
        
        st.write("**Deskripsi Masalah**")
        kendala = st.text_area("Ceritakan kendala IT yang Anda alami secara detail...", height=100)
        
        foto_awal = st.file_uploader("📸 Lampirkan Foto Error (Opsional)", type=['png', 'jpg', 'jpeg'])
        
        submitted = st.form_submit_button("🚀 Submit & Hubungkan ke AI", type="primary")
        
        if submitted:
            if kendala == "":
                st.warning("⚠️ Deskripsi Kendala wajib diisi!")
            else:
                nama_file_foto = "Tidak ada lampiran"
                if foto_awal:
                    os.makedirs("attachments", exist_ok=True)
                    nama_file_foto = f"{st.session_state.ticket_id}_{foto_awal.name}"
                    with open(f"attachments/{nama_file_foto}", "wb") as f: f.write(foto_awal.getbuffer())
                
                st.session_state.user_data['kendala_awal'] = kendala
                st.session_state.user_data['lampiran_awal'] = nama_file_foto
                st.session_state.first_prompt_pending = True
                st.session_state.ticket_status = "Open ⏳"
                st.session_state.page = "CHAT_CONSOLE"
                st.rerun()
    
    if st.button("⬅️ Kembali ke Beranda"): 
        st.session_state.page = "HOME"
        st.rerun()

# ==========================================
# ROUTER 5: CHAT CONSOLE
# ==========================================
elif st.session_state.page == "CHAT_CONSOLE":
    st.title("💡 Intelligent Service Desk")
    st.caption(f"ID Tiket: {st.session_state.ticket_id}")
    
    col_chat, col_info = st.columns([2.5, 1])

    with col_info:
        with st.container(border=True):
            st.subheader("⚙️ Status Tiket")
            waktu_wib = datetime.utcnow() + timedelta(hours=7)
            st.info(f"**Status:** {st.session_state.ticket_status}\n\n**Prioritas:** {st.session_state.ticket_priority}\n\n**Tanggal:** {waktu_wib.strftime('%d %b %Y')}")
            
            # TOMBOL SELESAI DIHILANGKAN DARI USER BIASA (User hanya bisa Chat & Lapor)
            st.caption("Status tiket hanya dapat diubah oleh divisi IT Support.")
                
            if st.button("⬅️ Jeda & Kembali ke Beranda", use_container_width=True):
                st.session_state.page = "HOME"
                st.rerun()

    with col_chat:
        with st.expander("📸 Lampirkan Foto Tambahan"):
            foto_kendala = st.file_uploader("Pilih Foto", type=['png', 'jpg', 'jpeg'], key=f"upload_{st.session_state.uploader_key}")

        chat_container = st.container(height=450)
        
        with chat_container:
            if st.session_state.first_prompt_pending:
                keluhan_awal = st.session_state.user_data['kendala_awal']
                lampiran_awal = st.session_state.user_data.get('lampiran_awal', 'Tidak ada lampiran')
                img_path = f"attachments/{lampiran_awal}" if lampiran_awal != "Tidak ada lampiran" else None
                    
                st.session_state.messages.append({"role": "user", "content": keluhan_awal, "image_path": img_path})
                
                if any(k in keluhan_awal.lower() for k in ['mati', 'rusak', 'terbakar', 'server']):
                    st.session_state.ticket_priority = "High 🔴"
                elif any(k in keluhan_awal.lower() for k in ['lupa', 'password', 'lambat']):
                    st.session_state.ticket_priority = "Medium 🟡"
                else:
                    st.session_state.ticket_priority = "Low 🟢"
                    
                st.session_state.ticket_status = "In Progress 🔄"
                
                if vector_db:
                    relevant_docs = vector_db.similarity_search(keluhan_awal, k=1)
                    konteks = relevant_docs[0].page_content
                else: konteks = "Panduan umum."
                
                # UPDATE PROMPT AI AGAR SANGAT SPESIFIK & NON-TEKNIS
                try:
                    client = Groq(api_key=api_key)
                    sys_prompt = f"""Sebagai AI Support Optima Resolve. Konteks: {konteks}.
                    ATURAN MUTLAK:
                    1. Berikan solusi spesifik secara LANGKAH DEMI LANGKAH (Gunakan angka 1, 2, 3).
                    2. Gunakan bahasa Indonesia yang SANGAT SEDERHANA agar mudah dipahami oleh orang awam/non-teknis. Jangan gunakan istilah rumit.
                    3. WAJIB TAMBAHKAN KALIMAT INI DI AKHIR JAWABAN: "Jika langkah di atas belum menyelesaikan masalah Anda, atau jika ini berkaitan dengan kerusakan fisik/server, silakan hubungi tim IT Support kami di Ext: 1122 atau email ke it.support@optima.com." """
                    
                    response = client.chat.completions.create(
                        model="llama-3.1-8b-instant",
                        messages=[{"role": "system", "content": sys_prompt}, {"role": "user", "content": keluhan_awal}]
                    )
                    full_response = response.choices[0].message.content
                except Exception:
                    full_response = "Maaf, mesin AI sedang offline. Silakan hubungi tim IT Support di Ext: 1122 atau email ke it.support@optima.com."
                
                st.session_state.messages.append({"role": "assistant", "content": full_response})
                st.session_state.first_prompt_pending = False
                
                st.session_state.waktu_awal = (datetime.utcnow() + timedelta(hours=7)).strftime("%Y-%m-%d %H:%M:%S")
                try:
                    db.collection('rekap_tiket').document(st.session_state.ticket_id).set({
                        'Waktu': st.session_state.waktu_awal, 'ID Tiket': st.session_state.ticket_id,
                        'Prioritas': st.session_state.ticket_priority, 'NIP': st.session_state.user_data['nip'],
                        'Nama': st.session_state.user_data['nama'], 'Divisi': st.session_state.user_data['divisi'],
                        'Telepon': st.session_state.user_data['telepon'], 'Kendala': f"USER:\n{keluhan_awal}",
                        'Solusi': f"AI:\n{full_response}", 'Lampiran': lampiran_awal, 'Status': st.session_state.ticket_status
                    })
                except: pass
                st.rerun()

            for message in st.session_state.messages:
                with st.chat_message(message["role"]):
                    if message.get("image_path"): st.image(message["image_path"], width=250)
                    st.write(message["content"])

        prompt_input = st.chat_input("Ketik balasan Anda di sini...")
        
        if prompt_input:
            filepath = None
            if foto_kendala:
                os.makedirs("attachments", exist_ok=True)
                nama_file_foto = f"{st.session_state.ticket_id}_{foto_kendala.name}"
                filepath = f"attachments/{nama_file_foto}"
                with open(filepath, "wb") as f: f.write(foto_kendala.getbuffer())
                st.session_state.uploader_key += 1
                
            st.session_state.messages.append({"role": "user", "content": prompt_input, "image_path": filepath})
            
            try:
                if vector_db:
                    relevant_docs = vector_db.similarity_search(prompt_input, k=1)
                    konteks = relevant_docs[0].page_content
                else: konteks = "Panduan dasar."
                    
                client = Groq(api_key=api_key)
                hist = "\n".join([f"{m['role']}: {m['content']}" for m in st.session_state.messages[-3:-1]])
                # UPDATE PROMPT AI LANJUTAN
                sys_prompt = f"""Sebagai AI Support. Konteks: {konteks}. History: {hist}. 
                Berikan panduan lanjutan langkah demi langkah (1, 2, 3) yang sangat sederhana untuk non-teknis. Jika mentok, arahkan ke IT Support Ext: 1122."""
                
                response = client.chat.completions.create(
                    model="llama-3.1-8b-instant",
                    messages=[{"role": "system", "content": sys_prompt}, {"role": "user", "content": prompt_input}]
                )
                full_response = response.choices[0].message.content
                st.session_state.messages.append({"role": "assistant", "content": full_response})
                
                semua_kendala = "\n\n".join([f"USER: {m['content']}" for m in st.session_state.messages if m["role"] == "user"])
                semua_solusi = "\n\n".join([f"AI: {m['content']}" for m in st.session_state.messages if m["role"] == "assistant"])
                db.collection('rekap_tiket').document(st.session_state.ticket_id).update({
                    'Kendala': semua_kendala, 'Solusi': semua_solusi
                })
            except Exception:
                st.error("Mesin AI Timeout.")
            
            st.rerun()

# ==========================================
# ROUTER 6: ADMIN DASHBOARD
# ==========================================
elif st.session_state.page == "ADMIN":
    st.title("🛠️ OPTIMA ANALYTICS")
    if st.button("Keluar dari Dashboard", type="primary"):
        logout()
    
    with st.container(border=True):
        try:
            tickets_ref = db.collection('rekap_tiket')
            docs = tickets_ref.stream()
            data_list = [doc.to_dict() for doc in docs]
                
            if len(data_list) > 0:
                df = pd.DataFrame(data_list)
                if 'Prioritas' not in df.columns: df['Prioritas'] = '-'
                if 'Status' not in df.columns: df['Status'] = 'Closed ✅'
                
                # MENGURUTKAN KOLOM DAN MENAMBAHKAN KOLOM 'No'
                df = df.sort_values(by='Waktu', ascending=False)
                df.insert(0, 'No', range(1, len(df) + 1))
                kolom_direksi = ['No', 'Waktu', 'ID Tiket', 'Prioritas', 'NIP', 'Nama', 'Divisi', 'Kendala', 'Solusi', 'Lampiran', 'Status']
                df = df[[c for c in kolom_direksi if c in df.columns]]
            else:
                df = pd.DataFrame(columns=['No', 'Waktu', 'ID Tiket', 'Prioritas', 'NIP', 'Nama', 'Divisi', 'Kendala', 'Solusi', 'Lampiran', 'Status'])

            if not df.empty:
                col_chart1, col_chart2 = st.columns(2)
                with col_chart1:
                    st.write("**📈 Tren Volume Tiket**")
                    df['Tanggal'] = pd.to_datetime(df['Waktu'], errors='coerce').dt.date
                    tiket_per_hari = df.groupby('Tanggal').size().reset_index(name='Jumlah')
                    if not tiket_per_hari.empty:
                        fig_line = px.line(tiket_per_hari, x='Tanggal', y='Jumlah', markers=True)
                        st.plotly_chart(fig_line, use_container_width=True)
                    
                with col_chart2:
                    st.write("**📊 Masalah per Divisi**")
                    div_count = df['Divisi'].value_counts().reset_index()
                    div_count.columns = ['Divisi', 'Jumlah']
                    fig_pie = px.pie(div_count, values='Jumlah', names='Divisi', hole=0.4)
                    st.plotly_chart(fig_pie, use_container_width=True)

                st.write("#### 📋 Database Log Enterprise")
                df_bersih = df.drop(columns=['Tanggal'], errors='ignore')
                st.dataframe(df_bersih, use_container_width=True, hide_index=True) 
                
                st.divider()
                st.write("#### 🔍 Inspektur Detail & Bukti Foto Tiket")
                list_tiket = df['ID Tiket'].dropna().unique()
                selected_ticket = st.selectbox("Pilih ID Tiket:", list_tiket)
                
                if selected_ticket:
                    tiket_row_awal = df[df['ID Tiket'] == selected_ticket].iloc[0]
                    nama_foto = tiket_row_awal.get('Lampiran', 'Tidak ada lampiran')
                    col_det1, col_det2 = st.columns([1.8, 1.2])
                    
                    with col_det1:
                        st.write(f"👤 **Pelapor:** {tiket_row_awal.get('Nama','')} — Status: **{tiket_row_awal.get('Status','')}**")
                        st.info(f"📌 **Kendala:**\n\n{tiket_row_awal.get('Kendala','')}")
                        st.success(f"🤖 **Solusi:**\n\n{tiket_row_awal.get('Solusi','')}")
                        
                    with col_det2:
                        st.write("🖼️ **Bukti Screenshot:**")
                        if nama_foto != "Tidak ada lampiran" and os.path.exists(f"attachments/{nama_foto}"):
                            st.image(f"attachments/{nama_foto}", use_container_width=True)
                        else: st.warning("⚠️ Tidak ada lampiran.")
                
                output = io.BytesIO()
                with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                    df_bersih.to_excel(writer, index=False, sheet_name='Data_Optima')
                st.download_button("📥 Download Excel Report", data=output.getvalue(), file_name=f"Report_{datetime.now().strftime('%Y%m%d')}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", type="primary")
            else:
                st.info("📊 Database Firebase Anda saat ini masih kosong.")
                
        except Exception as e: st.error(f"Error Database: {e}")

# WATERMARK PREMIUM
st.markdown("<br><br><div style='text-align: center; color: #64748B; font-size: 12px; padding: 30px; border-top: 1px solid #E2E8F0;'>© 2026 IT Division | Optima Resolve Enterprise Identity Management</div>", unsafe_allow_html=True)